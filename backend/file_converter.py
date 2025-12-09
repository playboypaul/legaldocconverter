import os
import tempfile
import subprocess
import logging
from pathlib import Path
from typing import Optional
import pypandoc
from docx import Document
import PyPDF2
from io import BytesIO

logger = logging.getLogger(__name__)

class FileConverter:
    """Handle file format conversions"""
    
    def __init__(self):
        # Use persistent storage directory
        storage_base = os.path.join(os.path.dirname(__file__), "storage")
        self.temp_dir = os.path.join(storage_base, "conversions")
        os.makedirs(self.temp_dir, exist_ok=True)
    
    async def convert_file(self, input_path: str, input_format: str, output_format: str, conversion_id: str) -> str:
        """Convert file from input format to output format"""
        try:
            output_filename = f"{conversion_id}_converted.{output_format}"
            output_path = os.path.join(self.temp_dir, output_filename)
            
            # Route to appropriate conversion method
            if input_format == "pdf":
                if output_format in ["txt", "html"]:
                    await self._convert_pdf_to_text_based(input_path, output_path, output_format)
                elif output_format == "docx":
                    await self._convert_pdf_to_docx(input_path, output_path)
                else:
                    await self._convert_with_pandoc(input_path, output_path, input_format, output_format)
            
            elif input_format in ["docx", "doc"]:
                if output_format == "pdf":
                    await self._convert_docx_to_pdf(input_path, output_path)
                elif output_format in ["txt", "html", "rtf", "odt"]:
                    await self._convert_with_pandoc(input_path, output_path, input_format, output_format)
                else:
                    await self._convert_docx_to_docx(input_path, output_path)
            
            elif input_format == "txt":
                await self._convert_text_based(input_path, output_path, input_format, output_format)
            
            else:
                # Use pandoc for other formats
                await self._convert_with_pandoc(input_path, output_path, input_format, output_format)
            
            if not os.path.exists(output_path):
                raise Exception(f"Conversion failed: Output file not created")
            
            logger.info(f"Successfully converted {input_format} to {output_format}")
            return output_path
            
        except Exception as e:
            logger.error(f"Conversion error: {str(e)}")
            raise Exception(f"Failed to convert file: {str(e)}")
    
    async def _convert_with_pandoc(self, input_path: str, output_path: str, input_format: str, output_format: str):
        """Convert using pandoc"""
        try:
            # Map formats to pandoc format names
            format_mapping = {
                "docx": "docx",
                "doc": "doc",
                "txt": "plain",
                "html": "html",
                "rtf": "rtf",
                "odt": "odt",
                "pdf": "pdf"
            }
            
            input_fmt = format_mapping.get(input_format, input_format)
            output_fmt = format_mapping.get(output_format, output_format)
            
            # Use pypandoc for conversion
            pypandoc.convert_file(
                input_path,
                output_fmt,
                outputfile=output_path,
                format=input_fmt
            )
            
        except Exception as e:
            # Fallback to command line pandoc
            try:
                cmd = [
                    "pandoc",
                    input_path,
                    "-f", input_format,
                    "-t", output_format,
                    "-o", output_path
                ]
                subprocess.run(cmd, check=True, capture_output=True, text=True)
            except subprocess.CalledProcessError as e:
                raise Exception(f"Pandoc conversion failed: {e.stderr}")
    
    async def _convert_pdf_to_text_based(self, input_path: str, output_path: str, output_format: str):
        """Convert PDF to text-based formats"""
        try:
            # Extract text from PDF
            with open(input_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            if output_format == "txt":
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
            elif output_format == "html":
                html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>Converted Document</title>
    <meta charset="UTF-8">
</head>
<body>
    <pre>{text}</pre>
</body>
</html>"""
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                    
        except Exception as e:
            raise Exception(f"PDF to text conversion failed: {str(e)}")
    
    async def _convert_pdf_to_docx(self, input_path: str, output_path: str):
        """Convert PDF to DOCX"""
        try:
            # Extract text from PDF
            with open(input_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            # Create DOCX document
            doc = Document()
            doc.add_paragraph(text)
            doc.save(output_path)
            
        except Exception as e:
            raise Exception(f"PDF to DOCX conversion failed: {str(e)}")
    
    async def _convert_docx_to_pdf(self, input_path: str, output_path: str):
        """Convert DOCX to PDF using pandoc"""
        try:
            await self._convert_with_pandoc(input_path, output_path, "docx", "pdf")
        except Exception as e:
            raise Exception(f"DOCX to PDF conversion failed: {str(e)}")
    
    async def _convert_docx_to_docx(self, input_path: str, output_path: str):
        """Copy DOCX file (for same format 'conversion')"""
        import shutil
        shutil.copy2(input_path, output_path)
    
    async def _convert_text_based(self, input_path: str, output_path: str, input_format: str, output_format: str):
        """Convert between text-based formats"""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if output_format == "html":
                html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>Converted Document</title>
    <meta charset="UTF-8">
</head>
<body>
    <pre>{content}</pre>
</body>
</html>"""
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
            
            elif output_format == "docx":
                doc = Document()
                doc.add_paragraph(content)
                doc.save(output_path)
            
            else:
                # For other formats, use pandoc
                await self._convert_with_pandoc(input_path, output_path, input_format, output_format)
                
        except Exception as e:
            raise Exception(f"Text conversion failed: {str(e)}")